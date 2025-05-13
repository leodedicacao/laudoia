# -----------------------------------------------------------------------------
# LAUDO GENERATOR GUI v1.7.4-debug - Adiciona Prints para Debug ORB
# -----------------------------------------------------------------------------
# Requisitos: Python 3.11, pip install opencv-python numpy docxtpl Pillow google-generativeai
# Necessário: 'modelo_laudo.docx', Chave API Gemini. DB ORB 'image_database.db'.
# Mudança v1.7.4-debug: Adicionados prints no terminal para depurar a saída
#                       da função encontrar_top_n_similares (ORB).
# -----------------------------------------------------------------------------

import cv2
import numpy as np
import sqlite3
import pickle
import os
import sys
import tkinter as tk
from tkinter import filedialog, simpledialog, scrolledtext, messagebox, Toplevel

# Imports para DOCX
try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Inches, Pt
except ImportError: messagebox.showerror("Erro Importação", "Instale: pip install docxtpl"); sys.exit(1)

# Imports para IA (Gemini) e Imagem
try:
    import google.generativeai as genai
    from PIL import Image, ImageTk
    llm_available = True
except ImportError: print("AVISO: 'google-generativeai' ou 'Pillow' não instalados."); llm_available = False

# --- Configurações Globais ---
DB_NAME = 'image_database.db'; TABLE_NAME = 'images'; db_exists = os.path.exists(DB_NAME)
orb = cv2.ORB_create(nfeatures=1000); bf = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=False)
N_TOP_RESULTS = 3; MIN_MATCH_COUNT = 10; RATIO_THRESH = 0.75
gemini_model_name = 'gemini-1.5-flash-latest'; llm_model = None; llm_model_active = False

# --- Funções DB e Imagem (ORB) ---
def inicializar_db():
    # (sem alterações)
    try:
        conn = sqlite3.connect(DB_NAME); cursor = conn.cursor()
        cursor.execute(f'CREATE TABLE IF NOT EXISTS {TABLE_NAME} (id INTEGER PRIMARY KEY AUTOINCREMENT, image_path TEXT UNIQUE NOT NULL, description TEXT NOT NULL, descriptors BLOB NOT NULL)')
        conn.commit(); conn.close(); print(f"DB '{DB_NAME}' OK."); return True
    except Exception as e: print(f"Erro DB: {e}"); messagebox.showerror("Erro DB", f"Erro inicializar DB: {e}"); return False

def calcular_descritores(caminho_imagem):
    # (sem alterações)
    try:
        img = cv2.imread(caminho_imagem, cv2.IMREAD_GRAYSCALE);
        if img is None: return None, None
        keypoints, descriptors = orb.detectAndCompute(img, None);
        if descriptors is not None and descriptors.shape[0] > 0: return keypoints, descriptors
        else: return keypoints, None
    except Exception as e: print(f"Erro processar {caminho_imagem}: {e}"); return None, None

def adicionar_imagem_db(caminho_imagem, descricao):
    # (sem alterações)
    if not os.path.exists(caminho_imagem): return f"Erro: Img não encontrada '{caminho_imagem}'"
    keypoints, descriptors = calcular_descritores(caminho_imagem)
    if descriptors is None or not isinstance(descriptors, np.ndarray) or descriptors.size == 0: return f"Erro: Não extraiu descritores válidos de '{os.path.basename(caminho_imagem)}'."
    try: descriptors_serializados = pickle.dumps(descriptors)
    except Exception as pickle_error: return f"Erro serializar '{os.path.basename(caminho_imagem)}': {pickle_error}"
    conn = sqlite3.connect(DB_NAME); cursor = conn.cursor(); status_message = ""
    try:
        cursor.execute(f'INSERT INTO {TABLE_NAME} (image_path, description, descriptors) VALUES (?, ?, ?)', (caminho_imagem, descricao, descriptors_serializados))
        conn.commit(); status_message = f"Sucesso: Img '{os.path.basename(caminho_imagem)}' adicionada/atualizada no DB ref. ORB."
    except sqlite3.IntegrityError: status_message = f"Info: Img '{os.path.basename(caminho_imagem)}' já existe no DB ref. ORB."
    except Exception as e: status_message = f"Erro adicionar DB ORB: {e}"
    finally: conn.close()
    return status_message

def encontrar_top_n_similares(caminho_imagem_consulta, top_n=N_TOP_RESULTS):
    # (sem alterações lógicas)
    results = []
    if not os.path.exists(caminho_imagem_consulta): return f"Erro: Consulta não encontrada '{caminho_imagem_consulta}'", []
    kp_consulta, desc_consulta = calcular_descritores(caminho_imagem_consulta)
    if desc_consulta is None or not isinstance(desc_consulta, np.ndarray) or desc_consulta.size == 0 or len(desc_consulta) < MIN_MATCH_COUNT : return f"Info: Desc insuf consulta '{os.path.basename(caminho_imagem_consulta)}'.", []
    conn = sqlite3.connect(DB_NAME); cursor = conn.cursor()
    try: cursor.execute(f"SELECT image_path, description, descriptors FROM {TABLE_NAME}"); all_images = cursor.fetchall()
    except Exception as e: conn.close(); return f"Erro ler DB ref: {e}", []
    finally: conn.close()
    if not all_images: return "Info: Nenhuma imagem no DB ref.", []
    for image_path_ref, description_ref, descriptors_serializados_ref in all_images:
        if os.path.abspath(caminho_imagem_consulta) == os.path.abspath(image_path_ref): continue
        try:
            desc_db = pickle.loads(descriptors_serializados_ref)
            if desc_db is None or not isinstance(desc_db, np.ndarray) or desc_db.size == 0 or len(desc_db) < MIN_MATCH_COUNT: continue
            if desc_consulta.dtype != desc_db.dtype or desc_consulta.shape[1] != desc_db.shape[1]: continue
            matches = bf.knnMatch(desc_consulta, desc_db, k=2)
            good_matches = []
            if matches and len(matches) > 0 and len(matches[0]) >= 2:
                 for m, n in matches:
                     if not isinstance(m, cv2.DMatch) or not isinstance(n, cv2.DMatch): continue
                     if m.distance < RATIO_THRESH * n.distance: good_matches.append(m)
            num_good_matches = len(good_matches)
            if num_good_matches >= MIN_MATCH_COUNT: results.append({'score': num_good_matches, 'path': image_path_ref, 'description': description_ref})
        except Exception as e: print(f"Aviso: Erro match ref '{os.path.basename(image_path_ref)}': {e}. Pulando.")
    results.sort(key=lambda x: x['score'], reverse=True)
    if not results: return f"Info: Nenhuma ref similar encontrada (limiar {MIN_MATCH_COUNT}).", []
    else: return f"Info: Encontrados {len(results)} resultados similares.", results[:top_n]

# --- Classe para a Janela de Edição ---
class EditDialog(Toplevel):
    # (Sem alterações desde v1.7.3)
    def __init__(self, parent, parent_app, image_path, image_name, initial_text):
        super().__init__(parent)
        self.transient(parent); self.title(f"Revisar Análise IA - {image_name}"); self.minsize(650, 450); self.resizable(True, True)
        self.parent_app = parent_app; self.result = None; self.image_path = image_path; self.initial_text = initial_text
        self.grid_rowconfigure(1, weight=1); self.grid_columnconfigure(0, weight=1)
        img_frame = tk.Frame(self, bd=1, relief=tk.SUNKEN); img_frame.grid(row=0, column=0, pady=(10, 5), padx=10, sticky='ew')
        self.img_label = tk.Label(img_frame); self.img_label.pack(pady=5); self.load_and_display_image()
        text_frame = tk.LabelFrame(self, text="Análise Técnica Gerada pela IA (Edite se necessário)"); text_frame.grid(row=1, column=0, pady=5, padx=10, sticky='nsew')
        text_frame.grid_rowconfigure(0, weight=1); text_frame.grid_columnconfigure(0, weight=1)
        self.text_area = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD, height=10, font=("Arial", 10)); self.text_area.grid(row=0, column=0, sticky='nsew', padx=5, pady=5)
        self.text_area.insert(tk.END, self.initial_text if self.initial_text else "[Erro ou Nenhuma análise gerada pela IA]")
        button_frame = tk.Frame(self); button_frame.grid(row=2, column=0, pady=(5, 10), padx=10, sticky='ew')
        button_frame.grid_columnconfigure(0, weight=1); button_frame.grid_columnconfigure(1, weight=1)
        confirm_button = tk.Button(button_frame, text="Confirmar e Adicionar ao Banco ORB", command=self.on_confirm, width=30, bg="lightgreen"); confirm_button.grid(row=0, column=0, padx=(10,5))
        skip_button = tk.Button(button_frame, text="Pular Análise desta Foto", command=self.on_skip, width=20); skip_button.grid(row=0, column=1, padx=(5, 10))
        self.protocol("WM_DELETE_WINDOW", self.on_skip); self.grab_set(); self.wait_window()
    def load_and_display_image(self):
        try:
            img = Image.open(self.image_path); max_width = 650; img_ratio = img.height / img.width; new_width = min(img.width, max_width); new_height = int(new_width * img_ratio); max_height = 300
            if new_height > max_height: new_height = max_height; new_width = int(new_height / img_ratio)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS); photo = ImageTk.PhotoImage(img); self.img_label.config(image=photo); self.img_label.image = photo
        except Exception as e: print(f"Erro carregar/exibir img diálogo: {e}"); self.img_label.config(text=f"[Erro carregar img: {os.path.basename(self.image_path)}]")
    def on_confirm(self):
        confirmed_text = self.text_area.get("1.0", tk.END).strip()
        if not confirmed_text: messagebox.showwarning("Texto Vazio", "Texto da análise vazio.", parent=self); return
        if self.parent_app: self.parent_app.add_confirmed_analysis_to_orb_db(self.image_path, confirmed_text)
        self.result = confirmed_text; self.destroy()
    def on_skip(self): self.result = None; self.destroy()

# --- Classe Principal da Aplicação GUI ---
class App:
    def __init__(self, root_element):
        # (igual à v1.7.2)
        self.root = root_element
        self.root.title("Gerador de Laudos v1.7.4-debug (LLM + ORB + Edição IA)") # Versão
        self.root.geometry("800x750")
        self.lista_fotos_pericia = []; self.template_path = ""; self.llm_model_active = False

        main_frame = tk.Frame(self.root); main_frame.pack(fill=tk.BOTH, expand=True)
        db_frame = tk.LabelFrame(main_frame, text="Banco de Imagens de Referência (ORB)"); db_frame.pack(padx=10, pady=5, fill="x")
        add_ref_button = tk.Button(db_frame, text="Adicionar Imagem de Referência", command=self.gui_adicionar_imagem_referencia, width=25); add_ref_button.pack(side=tk.LEFT, padx=5, pady=5)

        laudo_frame = tk.LabelFrame(main_frame, text="Dados do Laudo e Processo"); laudo_frame.pack(padx=10, pady=5, fill="x")
        laudo_frame.columnconfigure(1, weight=1); laudo_frame.columnconfigure(3, weight=1)
        row_counter = 0
        def add_laudo_field(parent, label_text, row, col, width=35, sticky="ew"):
            label = tk.Label(parent, text=label_text); label.grid(row=row, column=col*2, padx=(10,2), pady=3, sticky="w")
            entry = tk.Entry(parent, width=width); entry.grid(row=row, column=col*2 + 1, padx=(0,10), pady=3, sticky=sticky)
            return entry
        self.entry_local = add_laudo_field(laudo_frame, "Local:", row_counter, 0); row_counter += 1
        self.entry_data = add_laudo_field(laudo_frame, "Data Vistoria:", row_counter, 0, width=15, sticky="w"); row_counter += 1
        self.entry_cliente = add_laudo_field(laudo_frame, "Cliente/Solic.:", row_counter, 0); row_counter += 1
        self.entry_objetivo = add_laudo_field(laudo_frame, "Objetivo:", row_counter, 0); row_counter += 1
        self.entry_processo = add_laudo_field(laudo_frame, "Nº Processo:", row_counter, 0); row_counter += 1
        self.entry_vara = add_laudo_field(laudo_frame, "Vara:", row_counter, 0); row_counter += 1
        self.entry_forum = add_laudo_field(laudo_frame, "Fórum:", row_counter, 0); row_counter += 1
        self.entry_tipo_acao = add_laudo_field(laudo_frame, "Tipo de Ação:", row_counter, 0); row_counter += 1
        self.entry_autor = add_laudo_field(laudo_frame, "Autor(a):", row_counter, 0); row_counter += 1
        self.entry_reu = add_laudo_field(laudo_frame, "Réu(ré):", row_counter, 0)
        row_counter = 0
        self.entry_inicial = add_laudo_field(laudo_frame, "Ref. Inicial:", row_counter, 1); row_counter += 1
        self.entry_contestacao = add_laudo_field(laudo_frame, "Ref. Contestação:", row_counter, 1); row_counter += 1
        self.entry_decisao = add_laudo_field(laudo_frame, "Ref. Decisão:", row_counter, 1); row_counter += 1
        self.entry_participante1 = add_laudo_field(laudo_frame, "Participante 1:", row_counter, 1); row_counter += 1
        self.entry_participante2 = add_laudo_field(laudo_frame, "Participante 2:", row_counter, 1)

        files_frame = tk.LabelFrame(main_frame, text="Arquivos Necessários"); files_frame.pack(padx=10, pady=5, fill="x")
        template_subframe = tk.Frame(files_frame); template_subframe.pack(fill='x')
        template_button = tk.Button(template_subframe, text="Selecionar Template (.docx)", command=self.gui_selecionar_template, width=25); template_button.pack(side=tk.LEFT, padx=5, pady=5)
        self.label_template = tk.Label(template_subframe, text="Nenhum template.", anchor='w'); self.label_template.pack(side=tk.LEFT, padx=5, pady=5, fill='x', expand=True)
        fotos_subframe = tk.Frame(files_frame); fotos_subframe.pack(fill='x')
        fotos_button = tk.Button(fotos_subframe, text="Selecionar Fotos da Perícia", command=self.gui_selecionar_fotos_pericia, width=25); fotos_button.pack(side=tk.LEFT, padx=5, pady=5)
        self.label_fotos_pericia = tk.Label(fotos_subframe, text="Nenhuma foto.", anchor='w'); self.label_fotos_pericia.pack(side=tk.LEFT, padx=5, pady=5, fill='x', expand=True)

        generate_button = tk.Button(main_frame, text=">> GERAR LAUDO <<", command=self.gui_gerar_laudo_completo, font=("Arial", 12, "bold"), height=2, bg="lightgreen"); generate_button.pack(padx=10, pady=10, fill="x")

        status_frame = tk.LabelFrame(main_frame, text="Status / Log"); status_frame.pack(padx=10, pady=5, fill="both", expand=True)
        self.status_text_area = scrolledtext.ScrolledText(status_frame, wrap=tk.WORD, state=tk.DISABLED, height=10, font=("Courier New", 9)); self.status_text_area.pack(padx=5, pady=5, fill="both", expand=True)

        if llm_available: self.configure_llm_api()
        if not db_exists:
            if inicializar_db(): self.update_status("Banco de dados ORB inicializado.")
            else: self.update_status("FALHA ao inicializar banco de dados ORB.")
        else: self.update_status("Aplicação pronta. Banco de dados ORB encontrado.")

    # --- Métodos da Classe App ---
    def configure_llm_api(self):
        # (sem alterações)
        global llm_model, llm_available, gemini_model_name
        try:
            api_key = os.environ.get("GEMINI_API_KEY")
            if not api_key: self.update_status("Chave API Gemini não encontrada. LLM desabilitado."); self.llm_model_active = False; llm_available = False; return
            else: self.update_status("Chave API Gemini encontrada.")
            genai.configure(api_key=api_key)
            llm_model = genai.GenerativeModel(gemini_model_name)
            self.update_status(f"Modelo LLM ({gemini_model_name}) configurado."); self.llm_model_active = True
        except Exception as api_err: messagebox.showerror("Erro API Gemini", f"Falha config API/Modelo: {api_err}"); self.update_status(f"Erro Config API Gemini: {api_err}"); self.llm_model_active = False; llm_available = False

    def get_multimodal_llm_analysis(self, foto_path, analysis_context):
        # (sem alterações lógicas)
        global llm_model
        if not self.llm_model_active or llm_model is None: return "[API LLM não configurada]"
        self.update_status(f"    Analisando '{os.path.basename(foto_path)}' com Gemini..."); self.root.update_idletasks()
        try:
            img = Image.open(foto_path)
            prompt_parts = [ f"""[... Prompt v1.6.2 ...]""", img ] # Prompt PT-BR
            prompt_parts = [ f"""Você é um assistente virtual especializado em engenharia civil [...]. **IMPORTANTE: A sua resposta DEVE ser inteiramente em português do Brasil.**""", img ] # Prompt resumido para caber
            generation_config = genai.types.GenerationConfig(temperature=0.3, max_output_tokens=1024)
            safety_settings=[{"category": c, "threshold": "BLOCK_MEDIUM_AND_ABOVE"} for c in ["HARM_CATEGORY_HARASSMENT", "HARM_CATEGORY_HATE_SPEECH", "HARM_CATEGORY_SEXUALLY_EXPLICIT", "HARM_CATEGORY_DANGEROUS_CONTENT"]]
            response = llm_model.generate_content(prompt_parts, generation_config=generation_config, safety_settings=safety_settings, stream=False)
            if response.text: return response.text.strip()
            else:
                try: reason = response.prompt_feedback.block_reason; reason_msg = f"Resposta bloqueada ({reason})"
                except Exception: reason_msg = "Resposta vazia."
                self.update_status(f"    {reason_msg} p/ '{os.path.basename(foto_path)}'"); return f"[Análise Técnica IA: {reason_msg}]"
        except FileNotFoundError: error_msg = f"Erro: Img não encontrada '{foto_path}'"; print(error_msg); self.update_status(error_msg); return f"[Análise Técnica IA: {error_msg}]"
        except Exception as llm_err: detailed_error = str(llm_err); error_msg = f"Erro API Gemini: {llm_err.__class__.__name__}"; print(f"Erro detalhado Gemini: {detailed_error}"); self.update_status(f"    {error_msg}"); return f"[Análise Técnica IA: Erro API: {error_msg}]"

    def update_status(self, message):
        # (sem alterações)
        self.status_text_area.config(state=tk.NORMAL); self.status_text_area.insert(tk.END, f"{message}\n"); self.status_text_area.see(tk.END); self.status_text_area.config(state=tk.DISABLED); print(message); self.root.update_idletasks()

    def gui_adicionar_imagem_referencia(self):
        # (sem alterações)
        filepath = filedialog.askopenfilename(title="Selecionar Imagem de REFERÊNCIA (ORB)", filetypes=[("Imagens", "*.jpg *.jpeg *.png *.bmp *.tiff"), ("Todos", "*.*")])
        if not filepath: self.update_status("Adição cancelada."); return
        description = simpledialog.askstring("Descrição Imagem Referência (ORB)", f"Descrição para:\n{os.path.basename(filepath)}")
        if not description: self.update_status("Adição cancelada."); return
        self.update_status(f"Processando '{os.path.basename(filepath)}' p/ DB ORB..."); status_message = adicionar_imagem_db(filepath, description); self.update_status(status_message)

    def add_confirmed_analysis_to_orb_db(self, image_path, description):
        # (sem alterações)
        if not image_path or not description: self.update_status("Erro interno: Faltou path/desc p/ add DB ORB."); return
        self.update_status(f"    Add/Atualizando '{os.path.basename(image_path)}' no DB ORB..."); self.root.update_idletasks()
        status = adicionar_imagem_db(image_path, description); self.update_status(f"    Resultado DB ORB: {status}")

    def gui_selecionar_template(self):
        # (sem alterações)
        filepath = filedialog.askopenfilename(title="Selecionar Template (.docx)", filetypes=[("Documento Word", "*.docx")])
        if filepath: self.template_path = filepath; self.label_template.config(text=f"T: {os.path.basename(filepath)}"); self.update_status(f"Template: {filepath}")
        else: self.template_path = ""; self.label_template.config(text="Nenhum template.")

    def gui_selecionar_fotos_pericia(self):
        # (sem alterações)
        filepaths = filedialog.askopenfilenames(title="Selecionar Fotos da Perícia", filetypes=[("Imagens", "*.jpg *.jpeg *.png *.bmp *.tiff"), ("Todos", "*.*")])
        if filepaths: self.lista_fotos_pericia = sorted(list(filepaths)); self.label_fotos_pericia.config(text=f"{len(self.lista_fotos_pericia)} foto(s)."); self.update_status(f"{len(self.lista_fotos_pericia)} fotos selecionadas.")
        else: self.update_status("Seleção fotos cancelada.")

    # ***** MÉTODO PRINCIPAL ATUALIZADO (v1.7.4-debug) *****
    def gui_gerar_laudo_completo(self):
        if not self.template_path or not os.path.exists(self.template_path): messagebox.showerror("Erro Template", f"Template inválido."); return
        if not self.lista_fotos_pericia: messagebox.showerror("Erro Fotos", "Nenhuma foto selecionada."); return
        self.update_status("---------------------------------------------"); self.update_status("Coletando dados..."); self.root.update_idletasks()
        dados_laudo = {k: getattr(self, f"entry_{k}").get().strip() for k in ['local', 'data', 'cliente', 'objetivo', 'processo', 'vara', 'forum', 'tipo_acao', 'autor', 'reu', 'inicial', 'contestacao', 'decisao', 'participante1', 'participante2']}
        keys_map_context = {'processo': 'numero_processo', 'autor': 'nome_autor', 'reu': 'nome_reu', 'inicial': 'ref_inicial', 'contestacao': 'ref_contestacao', 'decisao': 'ref_decisao', 'participante1': 'participante1', 'participante2': 'participante2', 'tipo_acao': 'tipo_acao', 'vara': 'vara_processo', 'forum': 'forum_processo', 'cliente': 'nome_cliente', 'objetivo': 'objetivo_pericia', 'local': 'local_diligencia', 'data': 'data_vistoria'}
        analysis_context_for_llm = {keys_map_context.get(k, k): v for k, v in {k: v if v else "NI" for k, v in dados_laudo.items()}.items()}

        if llm_available and not self.llm_model_active:
             self.update_status("Tentando configurar API LLM..."); self.configure_llm_api();
             if not self.llm_model_active: self.update_status("AVISO: Análise LLM indisponível.")

        final_photo_data = []; self.update_status(f"Iniciando revisão interativa de {len(self.lista_fotos_pericia)} fotos...")
        for i, foto_path in enumerate(self.lista_fotos_pericia):
            nome_foto = os.path.basename(foto_path); self.update_status(f"--> Preparando Foto {i+1}/{len(self.lista_fotos_pericia)}: '{nome_foto}' p/ revisão...")
            initial_llm_text = "[Análise LLM não executada ou falhou]"
            if llm_available and self.llm_model_active: initial_llm_text = self.get_multimodal_llm_analysis(foto_path, analysis_context_for_llm)
            elif llm_available: initial_llm_text = "[Análise LLM: API não configurada]"
            dialog = EditDialog(self.root, self, foto_path, nome_foto, initial_llm_text); final_text_for_photo = dialog.result
            if final_text_for_photo is not None: self.update_status(f"    Análise '{nome_foto}' confirmada."); final_photo_data.append({'path': foto_path, 'name': nome_foto, 'llm_analysis': final_text_for_photo})
            else: self.update_status(f"    Análise '{nome_foto}' pulada.")

        if not final_photo_data: messagebox.showinfo("Cancelado", "Nenhuma análise confirmada."); self.update_status("Geração cancelada."); return
        self.update_status("Revisão concluída. Montando documento final...")
        try:
            doc = DocxTemplate(self.template_path); subdoc_fotos = doc.new_subdoc()
            all_ops_successful = True
            for i, photo_info in enumerate(final_photo_data):
                foto_path = photo_info['path']; nome_foto = photo_info['name']; final_llm_text = photo_info['llm_analysis']
                self.update_status(f"--> Adicionando Foto {i+1}/{len(final_photo_data)}: '{nome_foto}' ao laudo...")
                subdoc_fotos.add_paragraph(f"FOTO {i+1}: {nome_foto}")
                try: subdoc_fotos.add_picture(foto_path, width=Inches(5.0)); subdoc_fotos.add_paragraph()
                except Exception as img_err: error_msg=f"Erro add foto '{nome_foto}': {img_err}"; print(error_msg); self.update_status(error_msg); subdoc_fotos.add_paragraph(f"[[ERRO img '{nome_foto}']]"); all_ops_successful=False
                # Adiciona Análise LLM (Revisada)
                subdoc_fotos.add_paragraph(); p_llm = subdoc_fotos.add_paragraph(); p_llm.add_run("Análise Técnica Preliminar (IA - Revisada):\n").bold = True
                for line in final_llm_text.split('\n'): subdoc_fotos.add_paragraph(line)
                subdoc_fotos.add_paragraph()
                # Adiciona Análise Comparativa ORB
                subdoc_fotos.add_paragraph("Análise Comparativa (Banco de Referência ORB):")
                try:
                    status_msg, top_matches = encontrar_top_n_similares(foto_path, N_TOP_RESULTS)
                    # ***** INÍCIO DEBUG *****
                    print(f"DEBUG ORB para '{nome_foto}': Status='{status_msg}', Matches Encontrados={len(top_matches) if top_matches else 0}")
                    if top_matches:
                        print(f"DEBUG ORB Detalhes Matches para '{nome_foto}':")
                        for idx, match_debug in enumerate(top_matches):
                            print(f"  - Match {idx+1}: Score={match_debug.get('score')}, Path={match_debug.get('path')}, Desc={match_debug.get('description')}")
                    # ***** FIM DEBUG *****
                    if top_matches:
                        subdoc_fotos.add_paragraph("  Achados Similares:")
                        for match in top_matches:
                            ref_path = match.get('path', 'N/A'); ref_desc = match.get('description', 'N/A'); ref_score = match.get('score', 'N/A'); ref_nome_base = os.path.basename(ref_path)
                            p = subdoc_fotos.add_paragraph(); p.add_run(f"  - Ref: ").bold = True; p.add_run(f"'{ref_nome_base}' (Score: {ref_score})\n"); p.add_run(f"    Desc (Ref): ").italic = True; p.add_run(f"{ref_desc}")
                            try: subdoc_fotos.add_picture(ref_path, width=Inches(3.0)); subdoc_fotos.add_paragraph()
                            except Exception as img_ref_err: error_msg=f"Erro add foto ref '{ref_nome_base}': {img_ref_err}"; print(error_msg); self.update_status(error_msg); subdoc_fotos.add_paragraph(f"    [[ERRO img ref '{ref_nome_base}']]"); all_ops_successful=False
                    else: subdoc_fotos.add_paragraph(f"  - {status_msg}") # Mostra msg tipo "Nenhuma ref encontrada"
                except Exception as analyse_err: error_msg=f"Erro análise ORB '{nome_foto}': {analyse_err}"; print(error_msg); self.update_status(error_msg); subdoc_fotos.add_paragraph(f"[[ERRO ANÁLISE ORB]]"); all_ops_successful=False
                subdoc_fotos.add_paragraph("\n---------------------------------------------------------\n") # Separador final

            if not all_ops_successful: self.update_status("AVISO: Erros na montagem.")
            self.update_status("Preparando contexto final..."); self.root.update_idletasks()
            textos_padrao = { 'intro': "...", 'metodologia': "...", 'consideracoes': "...", 'conclusao': "...", 'encerramento': "...", 'perito_nome': "Seu Nome", 'perito_titulo': "Eng", 'perito_crea': "CREA"}
            context = {
                **{k: v if v else "NI" for k, v in dados_laudo.items()}, 'secao_detalhada_fotos': subdoc_fotos,
                **{f'texto_{k}': v for k, v in textos_padrao.items() if k not in ['perito_nome', 'perito_titulo', 'perito_crea']},
                'nome_perito': textos_padrao.get('perito_nome', ''), 'titulo_perito': textos_padrao.get('perito_titulo', ''), 'crea_perito': textos_padrao.get('perito_crea', ''),
            }
            keys_map = {'processo': 'numero_processo', 'autor': 'nome_autor', 'reu': 'nome_reu', 'inicial': 'ref_inicial', 'contestacao': 'ref_contestacao', 'decisao': 'ref_decisao', 'participante1': 'participante1', 'participante2': 'participante2', 'tipo_acao': 'tipo_acao', 'vara': 'vara_processo', 'forum': 'forum_processo', 'cliente': 'nome_cliente', 'objetivo': 'objetivo_pericia', 'local': 'local_diligencia', 'data': 'data_vistoria'}
            context = {keys_map.get(k, k): v for k, v in context.items()} # Renomeia
            self.update_status("Gerando documento Word..."); self.root.update_idletasks()
            doc.render(context)
            safe_local = "".join(c if c.isalnum() else "_" for c in dados_laudo['local'])[:30]; safe_data = "".join(c if c.isalnum() else "_" for c in dados_laudo['data'])[:15]
            output_path_suggestion = f"Laudo_{safe_local}_{safe_data}_RevIA_ORB.docx"
            output_path = filedialog.asksaveasfilename(title="Salvar Laudo Como...", defaultextension=".docx", filetypes=[("Documento Word", "*.docx")], initialfile=output_path_suggestion)
            if output_path:
                try: doc.save(output_path); final_msg = f"SUCESSO! Laudo salvo:\n'{output_path}'"; self.update_status(final_msg); self.update_status("--- FIM ---"); messagebox.showinfo("Concluído", final_msg)
                except Exception as save_error: error_msg = f"Erro SALVAR:\n{save_error}"; print(error_msg); self.update_status(f"FALHA SALVAMENTO: {save_error}"); messagebox.showerror("Erro Salvar", error_msg)
            else: self.update_status("Geração cancelada no salvamento.")
        except Exception as general_error:
            error_msg = f"Erro GERAL geração:\n{general_error}"; print(f"Erro GERAL: {general_error}"); self.update_status(f"FALHA GERAL: {general_error}"); messagebox.showerror("Erro Crítico Geração", error_msg)

# --- Inicialização ---
if __name__ == "__main__":
    lib_check_ok = True
    try: import docxtpl; from docx.shared import Inches; from PIL import Image, ImageTk
    except ImportError: print("Erro: 'docxtpl' ou 'Pillow' não instalada(s)."); lib_check_ok = False
    try: import google.generativeai
    except ImportError: llm_available = False; print("AVISO: google.generativeai não instalado.")
    if not lib_check_ok: sys.exit(1)
    root = tk.Tk()
    app = App(root)
    root.mainloop()